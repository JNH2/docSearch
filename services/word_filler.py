from docx import Document   

def replace_in_paragraph(paragraph, data: dict):
    """
    Safely replace placeholders inside a paragraph
    """
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, value)


def fill_word_template(template_path: str, output_path: str, data: dict):
    """
    Fill {{Field Name}} placeholders in a Word document
    Supports paragraphs AND tables
    """
    doc = Document(template_path)

    
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, data)


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, data)

    
    doc.save(output_path)
